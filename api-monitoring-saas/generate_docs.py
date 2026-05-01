import os
import markdown2
import pdfkit
from docx import Document
import re

def convert_md_to_docx(md_path, docx_path):
    # Read Markdown
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()
    
    # Very basic manual parsing for DOCX since python-docx doesn't take MD directly
    # Using python-docx directly to build it structure by structure is safer than external libraries
    
    doc = Document()
    doc.add_heading('Software Requirements Specification', 0)
    
    lines = md_text.split("\n")
    for line in lines:
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            pass # Title already added
        elif line.startswith("- **"):
            # Simple list item with bold prefix
            clean_line = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # Extract bold text if possible
            match = re.search(r'\*\*(.*?)\*\*(.*)', clean_line)
            if match:
                p.add_run(match.group(1)).bold = True
                p.add_run(match.group(2))
            else:
                p.add_run(clean_line)
        elif line.startswith("**"):
            # Bold line
            clean_line = line.replace("**", "").strip()
            p = doc.add_paragraph()
            p.add_run(clean_line).bold = True
        elif line.strip() == "---":
            # Page break or just paragraph
            pass
        elif line.strip():
            # Standard paragraph
            clean_line = line.replace("**", "")
            doc.add_paragraph(clean_line.strip())

    doc.save(docx_path)
    print(f"Generated {docx_path}")

def convert_md_to_pdf(md_path, pdf_path):
    try:
        from fpdf import FPDF
        
        with open(md_path, "r", encoding="utf-8") as f:
            md_text = f.read()

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_font("Arial", "B", 16)
        pdf.cell(200, 10, "Software Requirements Specification", ln=True, align="C")
        pdf.ln(10)
        
        lines = md_text.split("\n")
        for line in lines:
            line = line.replace("**", "").replace("`", "")
            
            if line.startswith("### "):
                pdf.set_font("Arial", "B", 12)
                pdf.cell(200, 10, line[4:].strip(), ln=True)
            elif line.startswith("## "):
                pdf.set_font("Arial", "B", 14)
                pdf.cell(200, 10, line[3:].strip(), ln=True)
            elif line.startswith("# "):
                pass
            elif line.startswith("- "):
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(0, 8, "  " + chr(149) + " " + line[2:].strip())
            elif line.strip() == "---":
                pdf.ln(5)
            elif line.strip():
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(0, 8, line.strip())
            else:
                pdf.ln(4)
                
        pdf.output(pdf_path)
        print(f"Generated {pdf_path}")
    except Exception as e:
        print("FPDF failed:", e)

if __name__ == "__main__":
    md_file = "SRS.md"
    docx_file = "SRS.docx"
    pdf_file = "SRS.pdf"
    
    convert_md_to_docx(md_file, docx_file)
    convert_md_to_pdf(md_file, pdf_file)
